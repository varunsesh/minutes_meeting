from openai import OpenAI
from dotenv import load_dotenv
from record_audio import AudioRecorder, RecordingApp
from datetime import datetime
import os



from docx import Document
  

def action_item_extraction(transcription):
    print("creating action items...")
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def key_points_extraction(transcription):
    print("key points extraction...")
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def abstract_summary_extraction(transcription):
    print("abstract summary rendering...")
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content

def meeting_minutes(transcription):
    print("Preparing meeting minutes...")
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    return {
        'meeting_date': formatted_date,
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }

def sentiment_analysis(transcription):
    print("generating sentiment")
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message.content
def save_as_docx(minutes, filename):
    # Check if the document exists and load it, otherwise create a new one
    if os.path.exists(filename):
        doc = Document(filename)
    else:
        doc = Document()

    # Iterate through the minutes in reverse order to add latest on top
    for key, value in reversed(minutes.items()):
        # Insert new content at the beginning of the document
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        # Insert a paragraph at the beginning
        p = doc.paragraphs[0].insert_paragraph_before(heading)
        p.style = doc.styles['Heading 1']
        # Insert text content
        p = doc.paragraphs[0].insert_paragraph_before(value)
        # Insert a line break after the section
        doc.paragraphs[0].insert_paragraph_before('')

    # Save the updated document
    doc.save(filename)


if __name__=="__main__":
    load_dotenv()

    client = OpenAI(
        # defaults to os.environ.get("OPENAI_API_KEY")
        # api_key="My API Key",
    )

    # Get the current date and time
    now = datetime.now()

# Format the date and time as "Month Day, Year"
    formatted_date = now.strftime("%B %d, %Y")
    recorder = AudioRecorder()
    app = RecordingApp(recorder)
    app.mainloop()
    transcription = app.get_transcript()
    minutes = meeting_minutes(transcription)
    print(minutes)
    save_as_docx(minutes, 'meeting_minutes.docx')
