import tkinter as tk
from tkinter import filedialog, scrolledtext
from threading import Thread
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import InstalledAppFlow
from io import BytesIO
import os
from docx import Document
from openai import OpenAI
from datetime import datetime
from openai import OpenAI 
from dotenv import load_dotenv

def authenticate_google_drive():
    creds = None

    # Load existing credentials from a file
    if os.path.exists('token.json'):
        creds = Credentials.from_authorized_user_file('token.json', SCOPES)

    # If there are no valid credentials available, or if they are expired, re-authenticate
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file('client_secrets.json', SCOPES)
            creds = flow.run_local_server(port=0)
        
        # Save the credentials for future use
        with open('token.json', 'w') as token:
            token.write(creds.to_json())

    return creds

def split_transcript(transcription, max_length=7000):
    segments = []
    if len(transcription) > 7500:
        while len(transcription) > 0:
            segment = transcription[:max_length].rsplit(' ', 1)[0]
            segments.append(segment)
            transcription = transcription[len(segment):].lstrip()
    else:
        segments.append(transcription)
    
    return segments

def process_segment(segment, is_final_segment):
    instruction = (
        "This is a part of a meeting transcript. Please read it carefully. "
        "If this is the final segment, provide a comprehensive summary, "
        "including action items and key points. "
        "Otherwise, remember the details for summarizing later."
    )
    
    prompt = f"{instruction}\n\n{segment}"

    if is_final_segment:
        prompt += "\n\nNow, provide the summary, action items, and key points based on the entire transcript, ensuring the summary does not exceed 3000 characters."

    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0.7,
        messages=[
            {
                "role": "system",
                "content": prompt
            }
        ]
    )
    
    summarized_content = response.choices[0].message.content
    return summarized_content
    

def process_and_summarize(segments):
    for i,segment in enumerate(segments):
        is_final_segment = (i == len(segments) - 1)
        processed_text = process_segment(segment, is_final_segment)

    return processed_text

def save_and_update_docx_on_drive(summary, local_filename, meeting_date):
    service = build('drive', 'v3', credentials=creds)

    request = service.files().export_media(fileId=google_doc_id, mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    fh = BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()
    
    fh.seek(0)
    document = Document(fh)

    document.add_paragraph("Meeting Minutes", style='Heading 1')
    document.add_paragraph(meeting_date, style='Normal')
    document.add_paragraph(summary, style='Normal')
    document.add_paragraph()

    temp_doc_path = "temp_" + local_filename
    document.save(temp_doc_path)

    # Re-upload the updated document
    file_metadata = {'name': local_filename, 'mimeType': 'application/vnd.google-apps.document'}
    media = MediaFileUpload(temp_doc_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', resumable=True)
    updated_file = service.files().update(fileId=google_doc_id, body=file_metadata, media_body=media, fields='id').execute()

    # Clean up the temporary local file
    os.remove(temp_doc_path)

    return updated_file.get('id')


def create_summary_audio(summary):
    response = client.audio.speech.create(
    model="tts-1",
    voice="onyx",
    input=summary
    )
    response.stream_to_file("meeting_summary.mp3")

load_dotenv()
client = OpenAI()
google_doc_id = os.environ.get("GOOGLE_DOC_ID")
SCOPES = ["https://www.googleapis.com/auth/drive"]
creds = authenticate_google_drive()


class TranscriptionApp(tk.Tk):
    def __init__(self, voice_recorder):
        super().__init__()
        self.voice_recorder = voice_recorder
        self.initialize_ui()

    def initialize_ui(self):
        self.title("Mr Minutes")
        self.geometry("750x500")
        self.start_recording_button = tk.Button(self, text="Start Recording", command=self.toggle_recording)
        self.start_recording_button.pack(pady=5)

        self.transcription_display = scrolledtext.ScrolledText(self, wrap=tk.WORD, height=10)
        self.transcription_display.pack(pady=10)
        self.transcription_display.config(state='disabled')

        self.initialize_additional_buttons()

    def initialize_additional_buttons(self):
        self.select_file_button = tk.Button(self, text="Upload Recording", command=self.load_audio_file)
        self.select_file_button.pack(pady=5)

        self.upload_transcript_button = tk.Button(self, text="Upload Transcript", command=self.load_transcript)
        self.upload_transcript_button.pack(pady=5)

        self.generate_transcript_button = tk.Button(self, text="Transcribe", command=self.create_transcript)
        self.generate_transcript_button.pack(pady=5)

        self.generate_minutes_button = tk.Button(self, text="Summarize", command=self.create_minutes)
        self.generate_minutes_button.pack(pady=5)

    def toggle_recording(self):
        if self.voice_recorder.is_active:
            self.end_recording()
            self.start_recording_button.config(text="Start Recording")
        else:
            self.begin_recording()
            self.start_recording_button.config(text="Stop Recording")

    def begin_recording(self):
        self.voice_recorder.begin_recording()
        self.recording_thread = Thread(target=self.voice_recorder.capture_audio, daemon=True)
        self.recording_thread.start()

    def end_recording(self):
        self.voice_recorder.end_recording()
        if self.recording_thread.is_alive():
            self.recording_thread.join()
        self.voice_recorder.transcription = self.voice_recorder.save_audio_and_transcribe()
        self.update_transcription_display(self.voice_recorder.transcription)

    def update_transcription_display(self, transcription):
        self.transcription_display.config(state='normal')
        self.transcription_display.delete(1.0, tk.END)
        self.transcription_display.insert(tk.INSERT, transcription)
        self.transcription_display.config(state='disabled')

    def load_audio_file(self):
        filepath = filedialog.askopenfilename(title="Select Audio File", filetypes=[("MP3 Files", "*.mp3"), ("WAV Files", "*.wav")])

        if filepath:
            self.voice_recorder.mp3_file_name = filepath  # Assuming VoiceRecorder can handle external file paths
            self.voice_recorder.transcription = self.voice_recorder._transcribe_audio()  # Call the transcription method
            self.update_transcription_display(self.voice_recorder.transcription)

    def load_transcript(self):
        filepath = filedialog.askopenfilename(title="Select Transcript File", filetypes=[("Text Files", "*.txt")])
        if filepath:
            with open(filepath, 'r', encoding='utf-8') as file:
                transcript = file.read()
                self.voice_recorder.transcription = transcript
            self.update_transcription_display(transcript)

    def create_transcript(self):
        if self.voice_recorder.mp3_file_name:
            self.voice_recorder.transcription = self.voice_recorder._transcribe_audio()
            self.update_transcription_display(self.voice_recorder.transcription)
        else:
            # Handle case where no audio file is loaded
            self.update_transcription_display("No audio file loaded for transcription.")

    def create_minutes(self):
        now = datetime.now()
        formatted_date = now.strftime("%B %d, %Y")
        transcription = self.voice_recorder.get_transcription()
        segments = split_transcript(transcription)
        summary = process_and_summarize(segments)
        doc_name = 'minutes_' + now.strftime("%B_%d_%Y") + '.docx'
        save_and_update_docx_on_drive(summary, doc_name, formatted_date)
        create_summary_audio(summary)
    