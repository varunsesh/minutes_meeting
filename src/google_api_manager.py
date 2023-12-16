import logging
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
import datetime
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from io import BytesIO
from PyQt5.QtWidgets import QInputDialog
import os, re
from docx import Document

class GoogleAPIManager:
    def __init__(self):
        """
        Initialize the Google API Manager.
        :return: None
        """
        self.logger = logging.getLogger(self.__class__.__name__)
        self.SCOPES = ['https://www.googleapis.com/auth/calendar.readonly', 'https://www.googleapis.com/auth/drive']
        self.credentials_file = os.path.join(os.path.dirname(__file__), 'config/creds.json')
        self.creds = None
        self.authenticate()
        self.doc_id = None

    def authenticate(self):
        """Authenticate with Google and set up the services.
        :return: None
        """
        self.creds = None
        token_path = os.path.join(os.path.dirname(__file__), 'config/token.json')
        if os.path.exists(token_path):
            self.creds = Credentials.from_authorized_user_file(token_path, self.SCOPES)

        if not self.creds or not self.creds.valid:
            
            if self.creds and self.creds.expired and self.creds.refresh_token:    
                self.logger.info('Refreshing token')
                self.creds.refresh(Request())
            
            else:
                flow = InstalledAppFlow.from_client_secrets_file(self.credentials_file, self.SCOPES)
                self.creds = flow.run_local_server(port=0)
            
            with open(token_path, 'w') as token:    
                token.write(self.creds.to_json())
                self.logger.info('Token generated successfully')

        self.calendar_service = build('calendar', 'v3', credentials=self.creds)
        self.drive_service = build('drive', 'v3', credentials=self.creds)
    
    def get_google_doc_id(self):
        """Get the Google Doc ID from the user or environment variable."""
        doc_id = os.environ.get("GOOGLE_DOC_ID")
        if not doc_id:
            doc_id = self.prompt_for_google_doc_id()
        self.logger.info(f"Google Docs Id: {doc_id}")
        return doc_id
    
    def prompt_for_google_doc_id(self):
        """Prompt the user to input the Google Doc ID."""
        text, ok = QInputDialog.getText(None, 'Google Docs ID', 'Enter Google Docs link or ID:')
        if ok and text:
            return self.extract_id_from_link(text)
        return None
    
    def extract_id_from_link(self, text):
        """Extract the ID from the Google Docs link."""
        match = re.search(r"/d/([a-zA-Z0-9-_]+)", text)
        return match.group(1) if match else text

    def download_document(self):
        """
        Download the Google Doc as a Word document.

        :return: The Word document
        """
        
        request = self.drive_service.files().export_media(fileId=self.doc_id, mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        fh = BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
            self.logger.info(f"Download status: {status}")
        
        fh.seek(0)
        return Document(fh)
    
    def update_docx_on_drive(self, summary, local_fname, date) -> str:
        """
        Update the Google Doc with the given summary, local file name, and date.

        :param summary: The summary to add to the Google Doc (list of strings)
        :param local_fname: The local temp file name of the Word document
        :param date: The date to add to the Google Doc
        :return: The updated Google Doc ID
        """

        self.doc_id = self.get_google_doc_id()

        if not self.doc_id:
            self.logger.info('Downloading copy of Google Doc')

            document = self.download_document()

            self.logger.info('Appending summary to Google Doc')
            document.add_paragraph("Meeting Notes", style='Heading 1')
            document.add_paragraph(date, style='Normal')
            
            document.add_paragraph("Objective", style='Heading 3')
            document.add_paragraph(summary["objective"], style ='Normal')

            document.add_paragraph("Key Points", style='Heading 3')
            for i, item in enumerate(summary["key_points"],1):
                document.add_paragraph(f'{i}. {item}', style='Normal')
            
            document.add_paragraph("Action Items", style='Heading 3')
            for i, item in enumerate(summary["action_items"],1):
                document.add_paragraph(f'{i}. {item}', style='Normal')
            
            temp_path = "temp_" + local_fname
            self.logger.info('Saving google document to temp file')
            document.save(temp_path)

            self.logger.info('Uploading google document to Cloud')
            updated_doc_id = self.upload_document(local_fname, temp_path)

            # Clean up the temporary local file
            os.remove(temp_path)

            return updated_doc_id    
        else: 
            self.logger.info('Failed to upload google document. document id missing')

    def upload_document(self, local_fname, temp_path):  
        """
        Upload the given local file to Google Drive as a new document.

        :param local_fname: The local file name of the Word document
        :param temp_path: The temporary path of the Word document
        :return: The Google Drive document ID
        """
        file_metadata = {'name': local_fname, 'mimeType': 'application/vnd.google-apps.document'}
        media = MediaFileUpload(temp_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', resumable=True)
        updated_file = self.drive_service.files().update(fileId=self.doc_id, body=file_metadata, media_body=media, fields='id').execute()
        
        return updated_file.get('id')
    
    def get_upcoming_meetings(self, max_results=10):
        """Fetch upcoming events from the user's calendar and look for attachments named 'minutes'."""
        now = datetime.datetime.utcnow().isoformat() + 'Z'  # 'Z' indicates UTC time
        events_result = self.calendar_service.events().list(
            calendarId='primary', 
            timeMin=now,
            maxResults=max_results, 
            singleEvents=True,
            orderBy='startTime'
        ).execute()
        events = events_result.get('items', [])

        meetings = []
        for event in events:
            start = event['start'].get('dateTime', event['start'].get('date'))
            meeting_link = event.get('hangoutLink', None)
            summary = event.get('summary', None)

            # Check for attachments that start with 'minutes'
            attachments = event.get('attachments', [])
            minutes_attachments = [att for att in attachments if att['title'].lower().startswith('minutes')]

            meetings.append({
                'start': start, 
                'meeting_link': meeting_link, 
                'summary': summary, 
                'minutes_attachments': minutes_attachments
            })

        return meetings
