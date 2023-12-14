import re
from docx import Document
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import InstalledAppFlow
from PyQt5.QtWidgets import QInputDialog
from io import BytesIO
import os

SCOPES = ["https://www.googleapis.com/auth/drive"]

class GoogleDocsManager:
    def __init__(self):
        self.creds = self.authenticate_google_drive()
        self.doc_id = self.get_google_doc_id()

    def authenticate_google_drive(self):
        creds = None

        # Load existing credentials from a file
        token_file_path = os.path.join(os.path.dirname(__file__), 'config/token.json')
        if os.path.exists(token_file_path):
            creds = Credentials.from_authorized_user_file(token_file_path, SCOPES)

        # If there are no valid credentials available, or if they are expired, re-authenticate
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                flow = InstalledAppFlow.from_client_secrets_file(os.path.join(os.path.dirname(__file__), 'config/creds.json'), SCOPES)
                creds = flow.run_local_server(port=0)
            
            # Save the credentials for future use
            with open(token_file_path, 'w') as token:
                token.write(creds.to_json())

        return creds
    def get_google_doc_id(self):
            doc_id = os.environ.get("GOOGLE_DOC_ID")
            if not doc_id:
                text, ok = QInputDialog.getText(None, 'Google Docs ID', 'Enter Google Docs link or ID:')
                if ok and text:
                    # Extract ID from link or use the provided ID
                    match = re.search(r"/d/([a-zA-Z0-9-_]+)", text)
                    if match:
                        doc_id = match.group(1)
                    else:
                        doc_id = text
            print("Google Docs Id", doc_id)
            return doc_id
    
    def download_document(self):
        service = build('drive', 'v3', credentials=self.creds)

        request = service.files().export_media(fileId=self.doc_id, mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        fh = BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
            print(status)
        
        fh.seek(0)
        return Document(fh)
    
    def save_and_update_docx_on_drive(self, summary, local_fname, date):
        document = self.download_document()
        document.add_paragraph("Meeting Minutes", style='Heading 1')
        document.add_paragraph(date, style='Normal')
        document.add_paragraph(summary, style='Normal')
        document.add_paragraph()
        temp_path = "temp_" + local_fname
        document.save(temp_path)
        updated_doc_id = self.upload_document(local_fname, temp_path)

        # Clean up the temporary local file
        os.remove(temp_path)

        return updated_doc_id
    

    def upload_document(self, local_fname, temp_path):
        service = build('drive', 'v3', credentials=self.creds)

        file_metadata = {'name': local_fname, 'mimeType': 'application/vnd.google-apps.document'}
        media = MediaFileUpload(temp_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', resumable=True)
        updated_file = service.files().update(fileId=self.doc_id, body=file_metadata, media_body=media, fields='id').execute()
        
        return updated_file.get('id')

